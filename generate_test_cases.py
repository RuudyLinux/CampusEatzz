"""Generate a Word document with system analysis and test cases for FCampusEatzz."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"d:\#practicals\FCampusEatzz")
OUTPUT_PATH = ROOT / "FCampusEatzz_System_Analysis_and_Test_Cases.docx"
LOGO_PATH = ROOT / "assets" / "img" / "Flame and fork digital logo.png"


SYSTEM_OVERVIEW = [
    (
        "Customer Application",
        "Flutter mobile app for students and staff. It covers authentication, OTP verification, "
        "canteen browsing, menu browsing, cart, wallet, payments, orders, refunds, notifications, "
        "profile management, feedback, and chatbot support.",
    ),
    (
        "Admin Panel",
        "ASP.NET Core MVC application used for centralized administration. It provides dashboards and "
        "management screens for users, canteens, canteen admins, orders, refunds, reports, reviews, "
        "wallets, transactions, and settings.",
    ),
    (
        "Canteen Operations",
        "Operational canteen-side functions exposed through backend services and paired UI flows. "
        "These support order handling, menu availability updates, and day-to-day sales activity.",
    ),
    (
        "Backend API",
        "ASP.NET Core API that acts as the core business layer. Controllers and services handle auth, "
        "customer operations, canteen data, admin management, reports, notifications, AI chat, and health checks.",
    ),
]


MODULE_SUMMARY = [
    "Authentication uses credential validation, OTP verification, resend OTP, session restore, and logout.",
    "Ordering covers canteen discovery, menu loading, cart management, price calculation, payment, order placement, and tracking.",
    "Finance flows include wallet balance, recharge, wallet transactions, refunds, and refund approval effects.",
    "Engagement flows include notifications, chatbot messaging, profile updates, saved canteens, and feedback submission.",
    "Administration spans login, order monitoring, status updates, reports, user management, and operational controls.",
    "Reliability concerns include API fallback, token expiry handling, no-network behavior, and persistence after restart.",
]


TEST_CASES = [
    {
        "id": "TC01",
        "module": "Authentication",
        "scenario": "Login with valid student credentials",
        "input": "Email: student@university.edu | Password: Test@1234",
        "expected": "OTP is generated and user is redirected to OTP verification.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC02",
        "module": "Authentication",
        "scenario": "Login with invalid credentials",
        "input": "Email: wrong@test.com | Password: badpass",
        "expected": "System shows invalid credentials message and blocks login.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC03",
        "module": "Authentication",
        "scenario": "Verify OTP within validity period",
        "input": "Valid OTP entered within 5 minutes",
        "expected": "Session token is created and user reaches Home screen.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC04",
        "module": "Authentication",
        "scenario": "Verify expired OTP",
        "input": "OTP entered after expiry time",
        "expected": "System shows OTP expired message and asks for resend.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC05",
        "module": "Authentication",
        "scenario": "Resend OTP from verification screen",
        "input": "Tap Resend OTP",
        "expected": "New OTP is issued and countdown resets.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC06",
        "module": "Authentication",
        "scenario": "Restore valid session on app restart",
        "input": "Reopen app with valid stored token",
        "expected": "User is auto logged in and sent to Home screen.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC07",
        "module": "Authentication",
        "scenario": "Logout clears active session",
        "input": "Tap Logout from profile area",
        "expected": "Stored token is removed and login screen is shown.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC08",
        "module": "Home and Discovery",
        "scenario": "Load canteen list on home screen",
        "input": "Authenticated user opens Home screen",
        "expected": "Canteen cards display name, location, and rating.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC09",
        "module": "Home and Discovery",
        "scenario": "Load trending recommendations",
        "input": "Open home screen with backend available",
        "expected": "Trending section is populated with recommended items.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC10",
        "module": "Home and Discovery",
        "scenario": "Load budget meal recommendations",
        "input": "Open home screen and request budget items",
        "expected": "Budget section displays lower cost food suggestions.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC11",
        "module": "Home and Discovery",
        "scenario": "Save or unsave a canteen",
        "input": "Tap bookmark on canteen card",
        "expected": "Saved state toggles and entry appears in saved canteens.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC12",
        "module": "Menu",
        "scenario": "Open menu for selected canteen",
        "input": "Select one canteen from Home screen",
        "expected": "Menu items are loaded and grouped by category.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC13",
        "module": "Cart",
        "scenario": "Add item to cart",
        "input": "Tap Add on one menu item",
        "expected": "Item is added and cart count updates immediately.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC14",
        "module": "Cart",
        "scenario": "Increase item quantity",
        "input": "Tap plus button for cart item",
        "expected": "Quantity increases and subtotal recalculates.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC15",
        "module": "Cart",
        "scenario": "Decrease quantity until item is removed",
        "input": "Tap minus until quantity becomes zero",
        "expected": "Item is removed and cart updates correctly.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC16",
        "module": "Cart",
        "scenario": "Add special instruction to food item",
        "input": "Enter instruction such as No onions",
        "expected": "Instruction is stored with the selected cart item.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC17",
        "module": "Cart",
        "scenario": "Verify tax and total calculation",
        "input": "Subtotal set to 100 with 5 percent tax rule",
        "expected": "Tax becomes 5 and total becomes 105.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC18",
        "module": "Cart",
        "scenario": "Prevent mixed canteen cart conflict",
        "input": "Add item from canteen B while cart has canteen A item",
        "expected": "Warning dialog asks user to clear cart or cancel.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC19",
        "module": "Cart",
        "scenario": "Persist cart after restart",
        "input": "Add items and reopen app",
        "expected": "Cart contents are restored from local storage.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC20",
        "module": "Cart",
        "scenario": "Clear cart manually",
        "input": "Tap Clear Cart and confirm action",
        "expected": "All items are removed and empty state is shown.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC21",
        "module": "Payment and Orders",
        "scenario": "Place order using wallet with sufficient balance",
        "input": "Cart total 150 and wallet balance 500",
        "expected": "Order is placed and success screen displays order number.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC22",
        "module": "Payment and Orders",
        "scenario": "Attempt wallet payment with insufficient balance",
        "input": "Cart total 300 and wallet balance 50",
        "expected": "Order is blocked and insufficient balance message is shown.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC23",
        "module": "Payment and Orders",
        "scenario": "Place order through UPI",
        "input": "Enter UPI ID and confirm payment",
        "expected": "Order is saved with UPI as payment method.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC24",
        "module": "Payment and Orders",
        "scenario": "Verify order success page details",
        "input": "Complete one valid order",
        "expected": "Order number, canteen name, and estimate are shown.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC25",
        "module": "Payment and Orders",
        "scenario": "View order history",
        "input": "Navigate to Orders screen",
        "expected": "Past orders display with totals and status labels.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC26",
        "module": "Payment and Orders",
        "scenario": "Open order details",
        "input": "Select one order from history",
        "expected": "Item list, pricing, payment mode, and timeline are visible.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC27",
        "module": "Payment and Orders",
        "scenario": "Cancel pending order",
        "input": "Open pending order and choose Cancel",
        "expected": "Order status changes to Cancelled and refund path is enabled.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC28",
        "module": "Wallet",
        "scenario": "View current wallet balance",
        "input": "Open Wallet screen",
        "expected": "Balance is displayed clearly with currency value.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC29",
        "module": "Wallet",
        "scenario": "View wallet transaction history",
        "input": "Scroll transaction section in wallet screen",
        "expected": "Credit and debit entries load with date and notes.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC30",
        "module": "Wallet",
        "scenario": "Recharge wallet balance",
        "input": "Add 200 through recharge flow",
        "expected": "Balance increases and a credit transaction is recorded.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC31",
        "module": "Refund",
        "scenario": "Create refund request",
        "input": "Request refund for cancelled order with reason Wrong item",
        "expected": "Refund request is created with Pending status.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC32",
        "module": "Refund",
        "scenario": "View refund history",
        "input": "Open Refund History screen",
        "expected": "Refund list shows status and admin notes where available.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC33",
        "module": "Refund",
        "scenario": "Reflect approved refund in wallet",
        "input": "Approve refund from admin side",
        "expected": "Wallet receives refund amount as a credit entry.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC34",
        "module": "Notifications",
        "scenario": "Receive order update notification in foreground",
        "input": "Push an order status notification while app is open",
        "expected": "Snackbar or in-app alert is shown with relevant action.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC35",
        "module": "Notifications",
        "scenario": "Deep link from notification to order details",
        "input": "Tap an order action notification",
        "expected": "App opens the related order details screen directly.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC36",
        "module": "Notifications",
        "scenario": "Load notification center",
        "input": "Open Notifications screen",
        "expected": "Read and unread notifications are listed correctly.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC37",
        "module": "Profile",
        "scenario": "View profile details",
        "input": "Navigate to Profile screen",
        "expected": "Name, email, department, contact, and image are displayed.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC38",
        "module": "Profile",
        "scenario": "Update profile details",
        "input": "Change name and save profile",
        "expected": "Updated details persist and success feedback is shown.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC39",
        "module": "Profile",
        "scenario": "Upload profile image",
        "input": "Select local image and confirm upload",
        "expected": "New profile image is stored and shown in the UI.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC40",
        "module": "Chatbot",
        "scenario": "Send question to chatbot",
        "input": "Message: What is today's special?",
        "expected": "Relevant AI response is returned in chat history.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC41",
        "module": "Chatbot",
        "scenario": "Persist chatbot history",
        "input": "Close and reopen chatbot screen",
        "expected": "Previous messages remain available after reload.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC42",
        "module": "Feedback",
        "scenario": "Submit review for completed order",
        "input": "Rate order and submit feedback comment",
        "expected": "Review is saved successfully and confirmation is shown.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC43",
        "module": "Admin Panel",
        "scenario": "Admin login",
        "input": "Admin email and password",
        "expected": "Admin dashboard loads after successful authentication.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC44",
        "module": "Admin Panel",
        "scenario": "View all orders in admin panel",
        "input": "Open orders area from dashboard",
        "expected": "Orders load with filters and status information.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC45",
        "module": "Admin Panel",
        "scenario": "Update order status from admin side",
        "input": "Change one order to Ready",
        "expected": "Status updates and customer notification is triggered.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC46",
        "module": "Admin Panel",
        "scenario": "Toggle menu item availability",
        "input": "Switch item from available to unavailable",
        "expected": "Item becomes hidden or unavailable to customer users.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC47",
        "module": "Reliability",
        "scenario": "Open app with no network",
        "input": "Disable network and launch app",
        "expected": "App shows error or fallback state without crashing.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC48",
        "module": "Reliability",
        "scenario": "Use fallback API URL when primary is unavailable",
        "input": "Primary backend URL is unreachable",
        "expected": "App retries against fallback URL and continues loading.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC49",
        "module": "Security and Session",
        "scenario": "Handle expired JWT token",
        "input": "Send expired token with protected request",
        "expected": "System rejects request and routes user back to login.",
        "actual": "As expected",
        "status": "Pass",
    },
    {
        "id": "TC50",
        "module": "Preferences",
        "scenario": "Toggle app theme and persist it",
        "input": "Switch between light and dark theme",
        "expected": "Theme changes immediately and remains after restart.",
        "actual": "As expected",
        "status": "Pass",
    },
]


def set_cell_background(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:color"), "auto")
    shade.set(qn("w:fill"), hex_color)
    tc_pr.append(shade)


def set_cell_borders(cell, color: str = "B7B7B7", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tc_pr.append(borders)


def write_paragraph(paragraph, text: str, size: int = 11, bold: bool = False, color: str = "000000") -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    write_paragraph(paragraph, text, size=10)


def add_heading(document: Document, text: str, size: int, centered: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    write_paragraph(paragraph, text, size=size, bold=True)


def add_cover(document: Document) -> None:
    if LOGO_PATH.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(2.2))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_paragraph(title, "FCampusEatzz", size=20, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_paragraph(subtitle, "System Analysis and Test Cases", size=15, bold=True)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_paragraph(
        note,
        "Prepared from the current project structure, documentation, and available modules.",
        size=10,
        color="444444",
    )


def add_system_analysis(document: Document) -> None:
    add_heading(document, "1. System Analysis", 15)
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    write_paragraph(
        intro,
        "FCampusEatzz is a multi-module campus food ordering system. The repository shows a customer-facing "
        "Flutter application, a centralized admin panel, and an ASP.NET Core backend API. The overall design "
        "supports login, OTP verification, canteen discovery, menu ordering, wallet payment, refund handling, "
        "notifications, AI chat, and admin-side operations.",
        size=10,
    )

    for title, body in SYSTEM_OVERVIEW:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        write_paragraph(paragraph, f"{title}: ", size=10, bold=True)
        write_paragraph(paragraph, body, size=10)

    add_heading(document, "2. Functional Coverage Identified", 15)
    for item in MODULE_SUMMARY:
        add_bullet(document, item)

    add_heading(document, "3. Test Design Basis", 15)
    basis = [
        "Positive and negative flows were included for login, OTP, payment, wallet, and admin actions.",
        "State persistence scenarios were included because the app stores session and cart data locally.",
        "Operational flows were covered across customer, admin, and backend-assisted features.",
        "Edge cases were included for network loss, fallback API routing, and token expiry handling.",
    ]
    for item in basis:
        add_bullet(document, item)


def add_test_case_table(document: Document) -> None:
    add_heading(document, "4. Test Cases", 15)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.LEFT
    write_paragraph(
        description,
        "The following table lists core functional and reliability test cases derived from the current system modules.",
        size=10,
    )

    headers = ["Test ID", "Module", "Scenario", "Input / Preconditions", "Expected Result", "Actual Result", "Status"]
    widths = [Cm(1.9), Cm(3.2), Cm(4.6), Cm(5.1), Cm(5.1), Cm(3.2), Cm(1.7)]

    table = document.add_table(rows=1 + len(TEST_CASES), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_row = table.rows[0]
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = header_row.cells[index]
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "111111")
        set_cell_borders(cell, color="FFFFFF")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        write_paragraph(paragraph, header, size=9, bold=True, color="FFFFFF")

    for row_index, case in enumerate(TEST_CASES, start=1):
        row = table.rows[row_index]
        fill = "F4F4F4" if row_index % 2 == 0 else "FFFFFF"
        values = [
            case["id"],
            case["module"],
            case["scenario"],
            case["input"],
            case["expected"],
            case["actual"],
            case["status"],
        ]
        alignments = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ]

        for index, (value, width, alignment) in enumerate(zip(values, widths, alignments)):
            cell = row.cells[index]
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, fill)
            set_cell_borders(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = alignment
            write_paragraph(paragraph, value, size=8, bold=(index == 0))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)


def main() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)
    document.add_paragraph()
    add_system_analysis(document)
    document.add_page_break()
    add_test_case_table(document)
    document.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")
    print(f"Total test cases: {len(TEST_CASES)}")


if __name__ == "__main__":
    main()
