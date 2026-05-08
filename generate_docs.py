from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

# ── Helper: set paragraph shading ─────────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'D0D0E0')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Style helpers ──────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x3c, 0x3c, 0x78)   # deep indigo
PRIMARY_L = RGBColor(0x55, 0x55, 0xa0)
ACCENT    = RGBColor(0x05, 0x96, 0x69)
DANGER    = RGBColor(0xe5, 0x3e, 0x3e)
DARK      = RGBColor(0x0f, 0x0f, 0x1e)
MUTED     = RGBColor(0x55, 0x55, 0x70)
WHITE     = RGBColor(0xff, 0xff, 0xff)

def heading1(text):
    p = doc.add_paragraph()
    shade_paragraph(p, '3c3c78')
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.1)
    run = p.add_run(f'  {text}')
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = WHITE
    return p

def heading2(text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'dde1f5')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.05)
    run = p.add_run(f' {text}')
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = PRIMARY
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = PRIMARY_L
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.3 + level*0.2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    return p

def kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'{key}: ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = PRIMARY_L
    r2 = p.add_run(value)
    r2.font.size = Pt(10); r2.font.color.rgb = DARK
    return p

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, '3c3c78')
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.add_row()
        bg = 'f8f8fc' if ri % 2 == 0 else 'ffffff'
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            shade_cell(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

# ─────────────────────────────────────────────────────────────────────────────
#  COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(cover, '3c3c78')
cover.paragraph_format.space_before = Pt(40)
cover.paragraph_format.space_after  = Pt(4)
r = cover.add_run('\nCampusEatzz\nSystem Documentation')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = WHITE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub, '3c3c78')
sub.paragraph_format.space_after = Pt(40)
rs = sub.add_run('University Campus Food Ordering Platform\n')
rs.font.size = Pt(12); rs.font.color.rgb = RGBColor(0xcc,0xcc,0xff)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(10)
rm = meta.add_run(f'Generated: {datetime.date.today().strftime("%B %d, %Y")}   |   Version 1.0   |   Confidential')
rm.font.size = Pt(9); rm.italic = True; rm.font.color.rgb = MUTED

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
#  1. SYSTEM OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────
heading1('1. System Overview')
body('CampusEatzz is a full-stack university campus food-ordering platform. It enables students and staff to browse canteen menus, place orders, pay via digital wallet, and track deliveries — all from a native mobile app. Canteen admins manage their menu and orders through a dedicated panel, while platform admins manage the entire system through a web dashboard.')

heading2('1.1 Architecture at a Glance')
make_table(
    ['Layer', 'Technology', 'Hosted On'],
    [
        ['Mobile App',       'Flutter (Dart)',             'Side-loaded APK / Play Store'],
        ['REST API Backend', 'ASP.NET Core + Dapper',      'Render (campuseatzz.onrender.com)'],
        ['Admin Dashboard',  'ASP.NET Core MVC + JS',      'Render (campuseatzz-admin.onrender.com)'],
        ['Database',         'MySQL 8',                    'Railway (managed cloud DB)'],
        ['Push Notifications','Firebase FCM',              'Google Cloud'],
        ['Email OTP',        'SMTP (Gmail) / Resend.dev',  'External API'],
        ['AI Features',      'OpenRouter (LLM)',           'External API'],
    ],
    [1.5, 1.8, 2.2]
)

heading2('1.2 High-Level Data Flow')
for step in [
    '1. Student opens Flutter app → authenticates via University ID + OTP (email)',
    '2. App calls REST API on Render → API queries Railway MySQL',
    '3. Student browses canteens/menu → adds items to cart',
    '4. Places order → wallet deducted → order stored in DB',
    '5. Push notification sent to canteen admin via Firebase FCM',
    '6. Canteen admin updates order status → customer notified',
    '7. Platform admin monitors everything via web dashboard',
]:
    bullet(step)

# ─────────────────────────────────────────────────────────────────────────────
#  2. PROJECT STRUCTURE
# ─────────────────────────────────────────────────────────────────────────────
heading1('2. Project Structure')
make_table(
    ['Folder / File', 'Purpose'],
    [
        ['backend/',          'ASP.NET Core REST API — all business logic, DB access, auth'],
        ['flutter_app/',      'Flutter native mobile app (student + canteen admin UI)'],
        ['admin_files/',      'ASP.NET Core MVC web admin dashboard'],
        ['JS/',               '15 JavaScript files powering admin panel interactivity'],
        ['CSS/style.css',     'Admin panel stylesheet (Liquid Glass design system)'],
        ['assets/',           'Static assets — logos, images'],
        ['universitycanteendb.sql', 'Full MySQL schema + seed data'],
    ],
    [2.0, 4.0]
)

# ─────────────────────────────────────────────────────────────────────────────
#  3. DATABASE SCHEMA
# ─────────────────────────────────────────────────────────────────────────────
heading1('3. Database Schema')
body('Database: MySQL 8 hosted on Railway. 17 tables.')

heading2('3.1 Table Reference')
make_table(
    ['Table', 'Key Columns', 'Purpose'],
    [
        ['users',            'id, UniversityId, first_name, last_name, email, role, status, password_hash, profile_image_url', 'All system users'],
        ['students',         'UniversityId (PK), course, semester, password_hash, email', 'Student-specific data'],
        ['university_staff', 'UniversityId (PK), department, DateOfBirth, password_hash, email', 'Staff-specific data'],
        ['canteens',         'id, name, description, image_url, status, display_order', 'Canteen outlets'],
        ['canteen_admins',   'id, canteen_id, username, password, name, email, contact, status, image_url', 'Canteen managers'],
        ['admin_users',      'id, name, email, password', 'Super-admin accounts'],
        ['menu_categories',  'id, name, description, display_order, is_active', 'Item categories'],
        ['menu_items',       'id, category_id, canteen_id, name, price, is_available, is_vegetarian, spice_level, is_deleted', 'Menu items (soft-delete)'],
        ['cart_items',       'CartItemId, UserId, CanteenId, MenuItemId, Quantity', 'Active cart state'],
        ['orders',           'id, user_id, canteen_id, order_number, order_type, total_amount, final_amount, payment_method, payment_status, order_status', 'Customer orders'],
        ['order_items',      'id, order_id, menu_item_id, item_name, quantity, unit_price, total_price', 'Order line items'],
        ['order_status_history', 'id, order_id, previous_status, new_status, changed_by', 'Status audit trail'],
        ['wallets',          'id, user_id, balance', 'User wallet balance'],
        ['wallet_transactions', 'id, user_id, transaction_id, amount, type, status, payment_gateway, order_id', 'All financial transactions'],
        ['reviews',          'id, user_id, canteen_id, order_id, rating, review_text, admin_response, status', 'User reviews'],
        ['contact_messages', 'id, name, email, subject, message, status, reply_message', 'Contact form submissions'],
        ['maintenance',      'id, maintenance_type (global/canteen), canteen_id, is_active, message, reason', 'Maintenance windows'],
        ['system_settings',  'id, setting_key, setting_value', 'App-wide config (tax, charges, hours)'],
        ['user_otps',        'id, identifier, otp_hash, expires_at, used_at', 'OTP session tracking'],
        ['auth_refresh_tokens', 'id, user_id, role, token_hash, expires_at_utc, revoked_at_utc', 'JWT refresh tokens'],
        ['app_notifications', 'id, title, body, type, canteen_id, reference_id', 'Notification records'],
        ['user_device_tokens','id, user_id, role, token, platform', 'FCM device tokens'],
    ],
    [1.5, 2.8, 2.2]
)

heading2('3.2 Key Enumerations')
make_table(
    ['Field', 'Allowed Values'],
    [
        ['users.role',               'student | staff | admin | canteen_admin'],
        ['users.status',             'active | inactive | banned'],
        ['orders.order_status',      'pending | confirmed | preparing | ready | completed | cancelled'],
        ['orders.payment_status',    'pending | paid | failed | refunded'],
        ['orders.order_type',        'dine_in | takeaway | delivery'],
        ['orders.payment_method',    'cash | card | upi | online | wallet'],
        ['wallet_transactions.type', 'credit | debit'],
        ['reviews.status',           'active | hidden'],
        ['contact_messages.status',  'unread | read | replied'],
        ['maintenance.maintenance_type', 'global | canteen'],
    ],
    [2.5, 4.0]
)

# ─────────────────────────────────────────────────────────────────────────────
#  4. BACKEND API
# ─────────────────────────────────────────────────────────────────────────────
heading1('4. Backend REST API')
kv('Base URL', 'https://campuseatzz.onrender.com')
kv('Framework', 'ASP.NET Core — Dapper ORM — MySQL connector')
kv('Auth', 'JWT Bearer — 12 h access token — 7 day refresh token')
kv('Email OTP', 'Gmail SMTP (primary) + Resend.dev (fallback)')
kv('Push', 'Firebase Cloud Messaging (FCM)')
kv('AI', 'OpenRouter API (chatbot + recommendations)')

heading2('4.1 Authentication Endpoints  /api/auth')
make_table(
    ['Method', 'Endpoint', 'Description', 'Auth'],
    [
        ['POST', '/api/auth/request-otp',    'Request OTP — verifies University ID + password', 'None'],
        ['POST', '/api/auth/resend-otp',     'Resend OTP to registered email',                  'None'],
        ['POST', '/api/auth/verify-otp',     'Verify OTP → returns JWT + session',              'None'],
        ['GET',  '/api/auth/me',             'Get current user profile',                        'JWT'],
        ['POST', '/api/admin/login',         'Super-admin email+password login → JWT',          'None'],
        ['POST', '/api/canteen-admin/login', 'Canteen admin email+password login → JWT',        'None'],
        ['POST', '/api/auth/refresh',        'Exchange refresh token for new access token',     'None'],
    ],
    [0.7, 2.2, 2.5, 0.7]
)

heading2('4.2 Customer Endpoints  /api/customer')
make_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET',  '/api/customer/wallet',                    'Get wallet balance'],
        ['GET',  '/api/customer/wallet/transactions',       'Wallet transaction history'],
        ['POST', '/api/customer/wallet/recharge',           'Recharge wallet'],
        ['GET',  '/api/customer/orders',                    'Order history'],
        ['POST', '/api/customer/orders',                    'Place new order'],
        ['POST', '/api/customer/orders/{ref}/cancel',       'Cancel order'],
        ['POST', '/api/customer/orders/{ref}/refund',       'Request refund'],
        ['POST', '/api/customer/reviews',                   'Submit review'],
        ['POST', '/api/customer/contact-messages',          'Submit contact message'],
        ['POST', '/api/customer/profile/upload-image',      'Upload profile photo'],
        ['GET',  '/api/customer/profile/image/{userId}',    'Serve profile photo (from DB BLOB)'],
    ],
    [0.7, 3.0, 2.8]
)

heading2('4.3 Canteen Admin Endpoints  /api/canteen')
make_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET',   '/api/canteen/dashboard',           'Dashboard stats + recent orders'],
        ['GET',   '/api/canteen/reports',             'Sales reports (summary, topItems, dailyTrend)'],
        ['GET',   '/api/canteen/menu-items',          'List menu items for canteen'],
        ['POST',  '/api/canteen/menu-items',          'Create menu item'],
        ['PUT',   '/api/canteen/menu-items/{id}',     'Update menu item'],
        ['DELETE','/api/canteen/menu-items/{id}',     'Soft-delete menu item'],
        ['PATCH', '/api/canteen/menu-items/{id}/availability', 'Toggle item availability'],
        ['GET',   '/api/canteen/orders',              'List orders (via operations controller)'],
        ['PATCH', '/api/canteen/orders/{id}/status',  'Update order status'],
        ['GET',   '/api/canteen/reviews',             'List reviews for canteen'],
        ['POST',  '/api/canteen/reviews/{id}/respond','Respond to review'],
        ['GET',   '/api/canteen/wallet',              'Canteen financial summary'],
        ['GET',   '/api/canteen/settings',            'Get canteen settings'],
        ['PUT',   '/api/canteen/settings/canteen',    'Update canteen info'],
        ['GET',   '/api/canteen/maintenance',         'Get maintenance status'],
        ['PUT',   '/api/canteen/maintenance/canteen', 'Set canteen maintenance mode'],
    ],
    [0.7, 3.0, 2.8]
)

heading2('4.4 Super Admin Endpoints  /api/admin')
make_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET',  '/api/admin/dashboard',              'Platform-wide stats'],
        ['GET',  '/api/admin/users',                  'All users (students + staff)'],
        ['POST', '/api/admin/users',                  'Create user'],
        ['POST', '/api/admin/users/bulk-import',      'Bulk import users (CSV/JSON)'],
        ['PUT',  '/api/admin/users/{id}',             'Edit user'],
        ['DELETE','/api/admin/users/{id}',            'Delete user'],
        ['GET',  '/api/admin/canteens',               'All canteens'],
        ['POST', '/api/admin/canteens',               'Create canteen'],
        ['PUT',  '/api/admin/canteens/{id}',          'Update canteen'],
        ['DELETE','/api/admin/canteens/{id}',         'Delete canteen'],
        ['GET',  '/api/admin/canteen-admins',         'All canteen admin accounts'],
        ['POST', '/api/admin/canteen-admins',         'Create canteen admin'],
        ['GET',  '/api/admin/contact-messages',       'All contact messages'],
        ['GET',  '/api/admin/reviews',                'All reviews'],
        ['GET',  '/api/admin/maintenance',            'Maintenance status (all)'],
        ['PUT',  '/api/admin/maintenance/system',     'Toggle global maintenance'],
        ['PUT',  '/api/admin/maintenance/canteen',    'Toggle per-canteen maintenance'],
    ],
    [0.7, 3.0, 2.8]
)

heading2('4.5 Public + Other Endpoints')
make_table(
    ['Method', 'Endpoint', 'Description'],
    [
        ['GET', '/api/public/canteens',       'Public canteen list (no auth)'],
        ['GET', '/api/public/settings',       'App settings (name, logo, hours)'],
        ['GET', '/api/maintenance/status',    'System maintenance flag (bypass)'],
        ['GET', '/api/health',                'DB health check'],
        ['POST','/api/chat/message',          'AI chatbot message'],
        ['GET', '/api/recommendations/personal', 'Personalized meal recommendations'],
        ['GET', '/api/notifications/history',  'User notification history'],
        ['POST','/api/notifications/device-token','Register FCM token'],
    ],
    [0.7, 2.8, 3.0]
)

# ─────────────────────────────────────────────────────────────────────────────
#  5. FLUTTER MOBILE APP
# ─────────────────────────────────────────────────────────────────────────────
heading1('5. Flutter Mobile App')
kv('Framework',  'Flutter 3.4+  (Dart)')
kv('State Mgmt', 'Provider (ChangeNotifier)')
kv('HTTP',       'Dio 5.x with multi-base-URL fallback + JWT interceptor')
kv('Storage',    'SharedPreferences (session, preferences)')
kv('Push',       'firebase_messaging 15.x')
kv('Images',     'cached_network_image — CDN/DB served')

heading2('5.1 Screens & Features')
make_table(
    ['Module', 'Screens', 'Key Functionality'],
    [
        ['Auth',          'login_screen, otp_screen, bootstrap_screen', 'University ID + password → OTP email → JWT'],
        ['Home',          'home_screen',            'Canteen grid, hero banner, search, filter, recommendations'],
        ['Menu',          'menu_screen',            'Category filter, vegetarian toggle, add-to-cart'],
        ['Cart',          'cart_screen',            'Review items, choose order type, checkout'],
        ['Orders',        'orders_screen, order_details_screen, order_success_screen', 'History, live status, invoice'],
        ['Wallet',        'wallet_screen',          'Balance, transaction history, recharge flow'],
        ['Profile',       'profile_screen, saved_canteens_screen', 'Profile photo upload, details, saved canteens'],
        ['Notifications', 'notifications_screen',  'In-app notification list, mark read'],
        ['Chat',          'chatbot_screen',         'AI chatbot for menu queries'],
        ['Payment',       'payment_screen',         'Wallet-based checkout confirmation'],
        ['Refunds',       'refund_request_screen, refund_history_screen', 'Submit and track refunds'],
        ['Contact',       'contact_us_screen',      'Submit support message'],
        ['Canteen Admin', 'canteen_admin_shell_screen', 'Dashboard, orders, menu, reports, reviews, wallet, settings'],
    ],
    [1.2, 2.0, 3.3]
)

heading2('5.2 State Providers')
make_table(
    ['Provider', 'Manages'],
    [
        ['AuthProvider',            'Login/logout, OTP flow, session, profile image update'],
        ['CanteenProvider',         'Canteen list, menu items, all-items cache for search'],
        ['CartProvider',            'Cart items, totals, clear on logout'],
        ['OrdersProvider',          'Order history, status polling'],
        ['WalletProvider',          'Balance, transaction history'],
        ['NotificationProvider',    'Unread count, notification list, FCM integration'],
        ['RecommendationProvider',  'AI-based item recommendations'],
        ['SavedCanteensProvider',   'Favourite canteens (local + remote)'],
        ['RefundProvider',          'Refund requests and status'],
        ['ChatProvider',            'Chat session, message history'],
        ['ThemeProvider',           'Dark / light mode toggle'],
    ],
    [2.2, 4.3]
)

heading2('5.3 Key Widgets')
make_table(
    ['Widget', 'Purpose'],
    [
        ['CustomerBottomNav',      'Animated liquid-pill bottom nav (Home / Wallet / Profile)'],
        ['NetworkFoodImage',       'CachedNetworkImage wrapper with shimmer placeholder + fallback'],
        ['AnimatedReveal',         'Crystallize entry animation (FadeTransition + ScaleTransition + slide)'],
        ['ShimmerLoader',          'Skeleton loading animation for lists and cards'],
        ['AppBackdrop',            'Liquid glass blob background (isolated RepaintBoundary)'],
        ['GradientHeader',         'Frosted glass sticky header used on all customer screens'],
        ['NotificationBellButton', 'Bell icon with unread badge, links to notifications screen'],
        ['AppStatusBadge',         'Color-coded order status pill (pending/confirmed/completed…)'],
        ['AppEmptyState',          'Consistent empty state placeholder with icon + action'],
        ['AppAsyncView',           'Loading/error/retry wrapper for async content'],
    ],
    [2.2, 4.3]
)

# ─────────────────────────────────────────────────────────────────────────────
#  6. ADMIN WEB DASHBOARD
# ─────────────────────────────────────────────────────────────────────────────
heading1('6. Admin Web Dashboard')
kv('URL',        'https://campuseatzz-admin.onrender.com')
kv('Framework',  'ASP.NET Core MVC — Razor views')
kv('Styling',    'Tailwind CSS + custom style.css (Liquid Glass theme)')
kv('JS',         '15 vanilla JavaScript files — no framework')

heading2('6.1 Pages')
make_table(
    ['Page', 'Route', 'Key Functions'],
    [
        ['Login',              '/Home/AdminLogin',              'Email + password auth → JWT stored in localStorage'],
        ['Dashboard',          '/Home/AdminDashboard',          'Platform KPIs, order counts, revenue, recent activity'],
        ['Orders',             '/Home/AdminAllOrders',          'Filter by status/canteen/date, view details, invoice print'],
        ['Users',              '/Home/AdminManageUsers',        'Create/edit/delete students + staff, bulk CSV import'],
        ['Wallets',            '/Home/AdminWallets',            'Per-user wallet balances'],
        ['Transactions',       '/Home/AdminWalletTransactions', 'Full financial transaction history with filters'],
        ['Refunds',            '/Home/AdminRefunds',            'Approve / reject refund requests'],
        ['Canteens',           '/Home/AdminManageCanteens',     'Create/edit/delete canteen outlets, upload images'],
        ['Canteen Admins',     '/Home/AdminManageCanteenAdmins','Create/edit canteen admin accounts'],
        ['Messages',           '/Home/AdminContactMessages',    'View and reply to contact messages'],
        ['Reviews',            '/Home/AdminReviews',            'Show/hide customer reviews'],
        ['Reports',            '/Home/AdminReports',            'Revenue analytics, top items, daily trends'],
        ['Settings',           '/Home/AdminSettings',           'App name, logo, tax, delivery charge, opening hours, maintenance'],
    ],
    [1.6, 2.4, 2.5]
)

heading2('6.2 JavaScript Files')
make_table(
    ['File', 'Responsibility'],
    [
        ['admin-common.js',               'API client, session check, JWT refresh, toast notifications (bottom-center pill)'],
        ['admin-login.js',                'Login form submission, error handling'],
        ['admin-dashboard.js',            'Load KPI cards, recent orders feed'],
        ['admin-manage-users.js',         'CRUD for users, bulk import, role toggle'],
        ['admin-manage-canteens.js',      'CRUD for canteens, image upload'],
        ['admin-manage-canteen-admins.js','CRUD for canteen admin accounts'],
        ['admin-all-orders.js',           'Filter/search orders, status update'],
        ['admin-order-invoice.js',        'Generate and print order invoice'],
        ['admin-wallets.js',              'Display wallet balances per user'],
        ['admin-wallet-transactions.js',  'Transaction log with export'],
        ['admin-refunds.js',              'Approve / reject refunds with reason'],
        ['admin-reviews.js',              'Moderate review visibility'],
        ['admin-contact-messages.js',     'View messages, compose replies'],
        ['admin-reports.js',              'Date-range analytics charts'],
        ['admin-settings.js',             'Platform settings form (maintenance, branding, pricing)'],
    ],
    [2.5, 4.0]
)

# ─────────────────────────────────────────────────────────────────────────────
#  7. CANTEEN ADMIN PANEL (Flutter)
# ─────────────────────────────────────────────────────────────────────────────
heading1('7. Canteen Admin Panel  (in Flutter App)')
body('Canteen admins log in directly inside the Flutter app using a separate login screen. After authentication they see a multi-tab shell screen.')

heading2('7.1 Tabs')
make_table(
    ['Tab', 'Index', 'Key Functions'],
    [
        ['Dashboard', '0', 'Pending / active / completed order counts, revenue, recent orders list'],
        ['Orders',    '1', 'Live order queue, update status (confirmed → preparing → ready → completed)'],
        ['Menu Items','2', 'Add / edit / delete menu items, toggle availability, upload image'],
        ['Reports',   '3', 'Date-range sales summary, top items, daily revenue trend'],
        ['Reviews',   '4', 'View customer ratings, compose official responses'],
        ['Wallet',    '5', 'Revenue breakdown by payment method, transaction log'],
        ['Settings',  '6', 'Canteen info, opening hours, profile, change password, maintenance mode'],
    ],
    [1.2, 0.6, 4.7]
)

# ─────────────────────────────────────────────────────────────────────────────
#  8. AUTHENTICATION & SECURITY
# ─────────────────────────────────────────────────────────────────────────────
heading1('8. Authentication & Security')

heading2('8.1 Customer Login Flow')
for s in [
    '1. User enters University ID + password in Flutter app',
    '2. POST /api/auth/request-otp → backend verifies password (BCrypt), sends 6-digit OTP to email',
    '3. OTP stored as BCrypt hash in user_otps table with 5-minute expiry',
    '4. User enters OTP → POST /api/auth/verify-otp → backend validates, clears OTP session',
    '5. Returns JWT (12 h) + refresh token (7 days) + user session DTO',
    '6. Flutter stores session in SharedPreferences; refreshes token silently on expiry',
]:
    bullet(s)

heading2('8.2 Admin / Canteen Admin Login Flow')
for s in [
    '1. Email + password submitted to /api/admin/login or /api/canteen-admin/login',
    '2. BCrypt hash verified against admin_users / canteen_admins table',
    '3. JWT returned — stored in localStorage (web admin) or SharedPreferences (Flutter)',
]:
    bullet(s)

heading2('8.3 Security Controls')
make_table(
    ['Control', 'Detail'],
    [
        ['Password hashing',      'BCrypt with cost factor 11-12 ($2a / $2y algorithms)'],
        ['Token expiry',          'Access token: 12 hours. Refresh token: 7 days with rotation'],
        ['OTP',                   '6-digit numeric, 5-minute TTL, single-use, BCrypt-hashed in DB'],
        ['Role-based access',     'student | staff | admin | canteen_admin — enforced per endpoint'],
        ['Maintenance bypass',    '/api/admin, /api/auth, /api/maintenance/status bypass global 503'],
        ['CORS',                  'Whitelist of allowed origins in appsettings.json'],
        ['Profile images',        'Stored as MEDIUMBLOB in MySQL (survives container redeploys)'],
        ['Refresh token revoke',  'revoked_at_utc set on logout; revoked tokens rejected on next use'],
    ],
    [2.0, 4.5]
)

# ─────────────────────────────────────────────────────────────────────────────
#  9. KEY INTEGRATIONS
# ─────────────────────────────────────────────────────────────────────────────
heading1('9. External Integrations')
make_table(
    ['Service', 'Purpose', 'Config Key'],
    [
        ['Firebase FCM',     'Push notifications to mobile app (order updates, maintenance alerts)', 'Fcm:ServerKey / Fcm:ProjectId'],
        ['Gmail SMTP',       'OTP email delivery (primary)',                                         'Smtp:Host / Smtp:Username / Smtp:Password'],
        ['Resend.dev',       'OTP email delivery (fallback when SMTP fails)',                        'Resend:ApiKey'],
        ['OpenRouter AI',    'AI chatbot + personalized food recommendations',                       'Ai:ApiKey / Ai:BaseUrl'],
        ['Railway MySQL',    'Managed cloud database',                                               'ConnectionStrings:DefaultConnection'],
        ['Render',           'Hosts API + Admin panel (auto-deploy from GitHub main branch)',        'GitHub push → auto deploy'],
    ],
    [1.4, 2.8, 2.3]
)

# ─────────────────────────────────────────────────────────────────────────────
#  10. DEPLOYMENT
# ─────────────────────────────────────────────────────────────────────────────
heading1('10. Deployment')

heading2('10.1 Backend API & Admin Panel (Render)')
for s in [
    'Both services auto-deploy from the main branch of the GitHub repository',
    'Push to main → Render detects change → build → deploy (zero-downtime swap)',
    'Free tier: spins down after 15 min inactivity; first request takes ~50 s cold start',
    'Environment variables (DB connection, JWT secret, email keys) set in Render dashboard',
]:
    bullet(s)

heading2('10.2 Database (Railway)')
for s in [
    'MySQL 8 — persistent managed DB — survives API redeployments',
    'Schema auto-applied on startup via EnsureCoreSchemaAsync() in Program.cs',
    'ALTER TABLE statements run idempotently (skip if column already exists)',
    'Profile images stored as MEDIUMBLOB — NOT filesystem — safe across redeploys',
]:
    bullet(s)

heading2('10.3 Flutter App')
for s in [
    'Primary base URL: https://campuseatzz.onrender.com',
    'Fallback URLs: http://10.0.2.2:5266 (Android emulator), http://localhost:5266',
    'Build: flutter build apk --release',
    'Distribute via direct APK or Play Store',
]:
    bullet(s)

heading2('10.4 Database Migration — maintenance table')
body('In the most recent deployment the legacy website_maintenance and maintenance_mode tables were merged into a single maintenance table:', indent=0.1)
make_table(
    ['Old tables (removed)', 'New table'],
    [
        ['website_maintenance  (global status only)',          'maintenance  (maintenance_type = \'global\', canteen_id = 0)'],
        ['maintenance_mode  (per-canteen, FK to canteens)',    'maintenance  (maintenance_type = \'canteen\', canteen_id = N)'],
    ],
    [3.0, 3.5]
)
body('Run this migration script on the live DB before deploying the updated backend:', indent=0.1)

# ─────────────────────────────────────────────────────────────────────────────
#  11. FEATURES SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
heading1('11. Feature Summary')

heading2('11.1 Customer (Student / Staff)')
features_customer = [
    'University ID + OTP (email) login',
    'Browse canteens with live maintenance/availability status',
    'Search across canteens and menu items',
    'Vegetarian / open-only filter',
    'Add to cart, choose dine-in / takeaway / delivery',
    'Wallet-based checkout',
    'Real-time order status tracking',
    'Order history and invoice view',
    'Refund request and tracking',
    'Star ratings and text reviews',
    'AI chatbot for menu and ordering help',
    'Personalized food recommendations (AI)',
    'Push notifications (order updates, canteen status)',
    'Dark mode / light mode',
    'Saved favourite canteens',
    'Profile photo upload (persisted in DB)',
    'Contact / support message submission',
]
for f in features_customer:
    bullet(f)

heading2('11.2 Canteen Admin')
for f in [
    'Dashboard: live order counts, revenue today, menu item count',
    'Order queue management with status progression',
    'Menu item CRUD with image upload, category, vegetarian flag, spice level',
    'Sales reports: summary stats, top items, daily revenue trend',
    'Read and respond to customer reviews',
    'Wallet overview: payment method breakdown, transaction history',
    'Settings: canteen info, opening hours, profile, password change',
    'Maintenance mode (close canteen for ordering)',
]:
    bullet(f)

heading2('11.3 Platform Admin (Web Dashboard)')
for f in [
    'Platform-wide KPI dashboard',
    'Full user management (create, edit, delete, bulk import)',
    'Canteen management (create, edit, upload logo)',
    'Canteen admin account management',
    'Global + per-canteen maintenance control',
    'Financial reports: revenue, transactions, refund approval',
    'Review moderation (hide/show)',
    'Contact message inbox with reply',
    'System settings (app name, logo, tax %, delivery charge, operating hours)',
    'Order invoice generation and print',
]:
    bullet(f)

# ─────────────────────────────────────────────────────────────────────────────
#  FOOTER / SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(end, '3c3c78')
end.paragraph_format.space_before = Pt(20)
end.paragraph_format.space_after  = Pt(20)
er = end.add_run('End of Document — CampusEatzz System Documentation')
er.font.size = Pt(10); er.font.color.rgb = RGBColor(0xcc,0xcc,0xff); er.italic = True

out = r'd:\#practicals\FCampusEatzz\CampusEatzz_System_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
